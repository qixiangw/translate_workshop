import streamlit as st
import time
import os
import tempfile
from docx import Document
import io

# 安装所需库：pip install streamlit python-docx

languages = {
    "English": "英语",
    "Deutsch": "德语",
    "Français": "法语",
    "日本語": "日语",
    "한국어": "韩语",
    "ไทย": "泰语",
    "Tiếng Việt": "越南语",
    "中文(简体)": "简体中文",
    "العربية": "阿拉伯语",
    "हिन्दी": "印地语",
    "Italiano": "意大利语",
    "Português": "葡萄牙语",
    "Русский": "俄语", 
    "Türkçe": "土耳其语",
    "Español": "墨西哥西班牙语",
    "Español (España)": "欧洲西班牙语",
    "中文(繁體)": "繁体中文"
}

def read_file(file_path, file_type):
    """读取文件内容"""
    if os.path.getsize(file_path) > 1000000:  # 1M = 1000000 bytes
        raise ValueError("File size exceeds 1M limit")
    
    if file_type == 'txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_type == 'docx':
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

# 保留之前的其他函数定义（split_text, identify_terms, translate, proofread, interpret, process_chunk）

def translate_document(file_path, term_list, file_type):
    content = read_file(file_path, file_type)
    chunks = split_text(content)
    translated_chunks = [process_chunk(chunk, term_list) for chunk in chunks]
    translated_document = "\n".join(translated_chunks)
    return translated_document

# Streamlit 应用
st.title("文档翻译应用")
st.write("离线文档翻译会对文档做自动分段，采用三步翻译工作流程完成文档的精翻；首先识别专有名词，然后进行初步翻译，接着识别可能存在的问题，并进行复译以提高质量。这个过程充分利用了Claude-3.5-sonnet模型的能力，在多个阶段进行智能处理。最后，工作流将所有翻译好的部分合并，形成一个完整的、高质量的翻译文档。")


st.subheader("选择目标语言")
target_lang = st.selectbox(
    "选择要翻译的目标语言",
    options=list(languages.keys()),
    format_func=lambda x: f"{x} - {languages[x]}"
    )
st.subheader("上传源文件")
# 上传文档
uploaded_file = st.file_uploader("上传要翻译的文档", type=["txt", "docx"])

# 上传专词列表
uploaded_terms = st.file_uploader("如有指定专有名词可以上传专词列表（每行一个词）", type=["txt"])


if uploaded_file:

    if uploaded_terms:
        # 读取专词列表
        term_list = [line.decode().strip() for line in uploaded_terms]
    else:
        term_list = []
    # 获取文件类型
    file_type = uploaded_file.name.split('.')[-1].lower()

    # 创建临时文件来存储上传的文档
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_file_path = tmp_file.name

    if st.button("开始翻译"):
        # 显示进度条
        progress_bar = st.progress(0)
        status_text = st.empty()

        # 模拟翻译过程
        for i in range(100):
            status_text.text(f"翻译进行中...{i+1}%")
            progress_bar.progress(i + 1)
            time.sleep(0.1)  # 模拟处理时间

        # 执行翻译
        try:
            translated_doc = translate_document(tmp_file_path, term_list, file_type)
            
            # 保存翻译结果
            if file_type == 'txt':
                # 保存为文本文件
                with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as output_file:
                    output_file.write(translated_doc)
                    output_path = output_file.name
                mime_type = "text/plain"
                output_filename = "translated_document.txt"
            else:
                # 保存为 Word 文档
                doc = Document()
                for para in translated_doc.split('\n'):
                    doc.add_paragraph(para)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_file:
                    doc.save(output_file.name)
                    output_path = output_file.name
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                output_filename = "translated_document.docx"

            # 提供下载链接
            with open(output_path, "rb") as file:
                btn = st.download_button(
                    label="下载翻译结果",
                    data=file,
                    file_name=output_filename,
                    mime=mime_type
                )

            status_text.text("翻译完成！")
        except Exception as e:
            st.error(f"翻译过程中出错: {str(e)}")
        finally:
            # 清理临时文件
            os.unlink(tmp_file_path)
            if 'output_path' in locals():
                os.unlink(output_path)

else:
    st.info("请上传文档和专词列表以开始翻译。")
